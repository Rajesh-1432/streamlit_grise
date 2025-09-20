import os
import glob
import zipfile
import streamlit as st
import pandas as pd
import tempfile
from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_core.prompts import PromptTemplate
from PyPDF2 import PdfReader
from docx import Document
from langchain.chains.question_answering import load_qa_chain
from dotenv import load_dotenv
from testscripts import script
import base64


# Initialize environment and API key
load_dotenv()
os.environ["OPENAI_API_KEY"] = os.getenv("OPENAI_API")


# === Enhanced UI Styling ===
st.markdown(
    """
    <style>
    /* Global background and text */
    .stApp {
        background: linear-gradient(135deg, #f5f7fa 0%, #c3cfe2 100%);
        font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
    }
    
    /* Main container */
    .main .block-container {
        padding-top: 2rem;
        padding-bottom: 2rem;
        max-width: 1000px;
    }

    /* Title styling */
    .center-title {
        text-align: center;
        font-size: 2.8em;
        font-weight: 700;
        color: #2c3e50;
        margin-bottom: 2rem;
        text-shadow: 0 2px 4px rgba(0,0,0,0.1);
        padding: 1rem 0;
        border-bottom: 3px solid #3498db;
        border-radius: 10px;
        background: linear-gradient(135deg, #ffffff 0%, #f8f9fa 100%);
        box-shadow: 0 4px 15px rgba(0,0,0,0.1);
    }

    /* Primary buttons */
    div.stButton > button {
        display: block;
        margin: 1rem auto;
        width: 300px !important;
        min-height: 50px;
        padding: 0.75rem 1.5rem;
        font-weight: 600;
        font-size: 1.1em;
        background: linear-gradient(135deg, #3498db 0%, #2980b9 100%);
        color: #ffffff;
        border: none;
        border-radius: 25px;
        transition: all 0.3s ease;
        box-shadow: 0 4px 15px rgba(52, 152, 219, 0.3);
        cursor: pointer;
    }

    div.stButton > button:hover {
        background: linear-gradient(135deg, #2980b9 0%, #1f639a 100%);
        transform: translateY(-2px);
        box-shadow: 0 6px 20px rgba(52, 152, 219, 0.4);
    }

    div.stButton > button:focus, 
    div.stButton > button:active {
        outline: none;
        transform: translateY(0px);
        box-shadow: 0 2px 10px rgba(52, 152, 219, 0.3);
    }

    /* Download button */
    .download-button {
        display: block;
        width: 350px;
        margin: 2rem auto;
        padding: 1rem 2rem;
        color: #ffffff;
        background: linear-gradient(135deg, #27ae60 0%, #2ecc71 100%);
        text-align: center;
        border-radius: 25px;
        text-decoration: none;
        font-size: 1.2em;
        font-weight: 600;
        transition: all 0.3s ease;
        box-shadow: 0 4px 15px rgba(39, 174, 96, 0.3);
        border: none;
    }

    .download-button:hover {
        background: linear-gradient(135deg, #229954 0%, #27ae60 100%);
        transform: translateY(-2px);
        box-shadow: 0 6px 20px rgba(39, 174, 96, 0.4);
        color: #ffffff;
        text-decoration: none;
    }

    /* Radio buttons container */
    .stRadio > div {
        background: #ffffff;
        padding: 1.5rem;
        border-radius: 15px;
        box-shadow: 0 2px 10px rgba(0,0,0,0.1);
        margin: 1rem 0;
    }

    .stRadio > div > label {
        font-weight: 600;
        color: #2c3e50;
        font-size: 1.1em;
        margin-bottom: 1rem;
        display: block;
    }

    /* Radio button options */
    .stRadio div[role="radiogroup"] > label {
        background: #f8f9fa;
        margin: 0.5rem 0;
        padding: 0.75rem 1rem;
        border-radius: 10px;
        border: 2px solid #e9ecef;
        transition: all 0.3s ease;
        cursor: pointer;
    }

    .stRadio div[role="radiogroup"] > label:hover {
        background: #e3f2fd;
        border-color: #3498db;
    }

    /* File uploader */
    section[data-testid="stFileUploader"] {
        border: 2px dashed #3498db;
        background: linear-gradient(135deg, #ffffff 0%, #f8f9fa 100%);
        border-radius: 15px;
        padding: 2rem;
        margin: 1rem 0;
        transition: all 0.3s ease;
    }

    section[data-testid="stFileUploader"]:hover {
        border-color: #2980b9;
        background: linear-gradient(135deg, #f8f9fa 0%, #e3f2fd 100%);
    }

    .uploadedFile {
        background: #ffffff;
        border-radius: 10px;
        padding: 0.5rem;
        margin: 0.25rem 0;
        box-shadow: 0 2px 5px rgba(0,0,0,0.1);
    }

    /* Success/Info/Warning messages */
    .stSuccess {
        background: linear-gradient(135deg, #d4edda 0%, #c3e6cb 100%);
        border: 1px solid #27ae60;
        border-radius: 10px;
        color: #155724;
    }

    .stInfo {
        background: linear-gradient(135deg, #cce7ff 0%, #b3daff 100%);
        border: 1px solid #3498db;
        border-radius: 10px;
        color: #0c5460;
    }

    .stWarning {
        background: linear-gradient(135deg, #fff3cd 0%, #ffeaa7 100%);
        border: 1px solid #f39c12;
        border-radius: 10px;
        color: #856404;
    }

    /* Progress indicators */
    .stProgress > div > div > div > div {
        background: linear-gradient(135deg, #3498db 0%, #2980b9 100%);
    }

    /* Scrolling headline improvements */
    .headline-container {
        background: linear-gradient(135deg, #ffffff 0%, #f8f9fa 100%);
        border-radius: 10px;
        padding: 1rem;
        margin: 1rem 0;
        box-shadow: 0 2px 10px rgba(0,0,0,0.1);
        border-left: 4px solid #3498db;
    }

    /* Sidebar styling */
    .css-1d391kg {
        background: linear-gradient(135deg, #ffffff 0%, #f8f9fa 100%);
    }

    /* Hide Streamlit menu and footer */
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    header {visibility: hidden;}
    
    /* Custom spacing */
    .custom-spacing {
        margin: 2rem 0;
    }
    
    /* File counter styling */
    .file-counter {
        background: linear-gradient(135deg, #e3f2fd 0%, #bbdefb 100%);
        border-radius: 20px;
        padding: 0.5rem 1rem;
        color: #1976d2;
        font-weight: 600;
        text-align: center;
        margin: 1rem 0;
        border: 1px solid #2196f3;
    }
    </style>
    """,
    unsafe_allow_html=True
)


# === Functions ===
def scrolling_headline(state, placeholder, headline_text: str):
    if state:
        # Enhanced scrolling effect
        html_code = f"""
        <div class="headline-container" style="overflow: hidden; white-space: nowrap; width: 100%; box-sizing: border-box;">
            <marquee scrollamount="4" behavior="scroll" direction="right" style="
                font-size: 1.1em;
                color: #2c3e50;
                font-weight: 600;
            ">
                🚀 {headline_text} ⚡
            </marquee>
        </div>
        """
        placeholder.markdown(html_code, unsafe_allow_html=True)
    else:
        placeholder.success(f"✅ {headline_text}")


def read_pdf(pdf_docs):
    text = ""
    for pdf in pdf_docs:
        pdf_reader = PdfReader(pdf)
        for page in pdf_reader.pages:
            text += page.extract_text()
    return text


def read_txt(txt_docs):
    text = ""
    for txt in txt_docs:
        text += txt.read().decode("utf-8")
    return text


def read_docx(docx_docs):
    text = ""
    for docx in docx_docs:
        doc = Document(docx)
        for para in doc.paragraphs:
            text += para.text + "\n"
    return text


def read_pdf_from_directory(directory):
    text = ""
    for filename in os.listdir(directory):
        if filename.endswith('.pdf'):
            pdf_path = os.path.join(directory, filename)
            pdf_reader = PdfReader(pdf_path)
            for page in pdf_reader.pages:
                text += page.extract_text()
    return text


def read_txt_from_directory(directory):
    text = ""
    for filename in os.listdir(directory):
        if filename.endswith('.txt'):
            txt_path = os.path.join(directory, filename)
            with open(txt_path, 'r', encoding='utf-8') as file:
                text += file.read() + "\n"
    return text


def read_docx_from_directory(directory):
    text = ""
    for filename in os.listdir(directory):
        if filename.endswith('.docx'):
            docx_path = os.path.join(directory, filename)
            doc = Document(docx_path)
            for para in doc.paragraphs:
                text += para.text + "\n"
    return text


def initialize_generative_model():
    return ChatOpenAI(model="gpt-4o", temperature=0.5)


def understand_tone_and_language(script_text):
    embeddings = OpenAIEmbeddings(model="text-embedding-3-large")
    vector_store = FAISS.from_texts([script_text], embedding=embeddings)
    vector_store.save_local("faiss_index")


def generate_script(user_input, generative_model):
    embeddings = OpenAIEmbeddings(model="text-embedding-3-large")
    new_db = FAISS.load_local("faiss_index", embeddings, allow_dangerous_deserialization=True)
    script_docs = new_db.similarity_search(user_input)

    prompt_template = """
    Take the persona of AI SAP Test Case Generator.
    A Testcase is an end to end workflow but not a single step process in SAP.
    Extract the overview/purpose of the input provided.
    For the overview extracted generate the test case/s.
    Do not generate test scripts or descriptions.
    Each test case should not exceed 10 words.
    Format: "Test Case X: [Description]"
    Provide at least 5 test cases.
    Context: \n{context}\n
    Generated Test Cases:
    """

    prompt = PromptTemplate(template=prompt_template, input_variables=["context"])
    chain = load_qa_chain(generative_model, chain_type="stuff", prompt=prompt)

    response = chain({"input_documents": script_docs}, return_only_outputs=True)
    return response["output_text"]


def test_cases_to_excel(test_cases, file_name="test_cases.xlsx"):
    data = []
    for case in test_cases.splitlines():
        if case.startswith("Test Case"):
            if ": " in case:
                description = case.split(": ", 1)[1].strip()
                test_case_number = case.split(": ", 1)[0].strip()
                data.append([test_case_number, description])
            else:
                data.append([case.strip(), ""])
    df = pd.DataFrame(data, columns=["Test Case Number", "Description"])
    return df


def convert_to_excel(test_scripts, file_path):
    with pd.ExcelWriter(file_path, engine='xlsxwriter') as writer:
        summary_sheet_name = "Test Cases"
        summary_data = []
        workbook = writer.book
        header_format = workbook.add_format({'bold': True, 'bg_color': '#3498db', 'font_color': '#ffffff'})
        title_format = workbook.add_format({'bold': True, 'bg_color': '#2c3e50', 'font_color': '#ffffff'})

        for i, (case_title, script_details) in enumerate(test_scripts.items(), start=1):
            steps = []
            for scenario, details in script_details.items():
                fiori_id = details.get("Fiori ID", "")
                step_desc = details.get("Step Description", {})
                exp_results = details.get("Expected Results", {})
                for step_num, step_desc_text in step_desc.items():
                    steps.append({
                        "Step Number": step_num,
                        "Step Description": step_desc_text,
                        "Expected Result": exp_results.get(step_num, ""),
                        "Fiori ID": fiori_id,
                    })

            df = pd.DataFrame(steps)
            sheet_name = f"TestScript {i}"
            df.to_excel(writer, index=False, sheet_name=sheet_name, startrow=1)
            worksheet = writer.sheets[sheet_name]
            worksheet.write(0, 0, case_title, title_format)
            summary_data.append([f"Test Case {i}: {case_title}"])
            for col_num, value in enumerate(df.columns.values):
                worksheet.write(1, col_num, value, header_format)

        summary_df = pd.DataFrame(summary_data, columns=["Test Case Overview"])
        summary_df.to_excel(writer, index=False, sheet_name=summary_sheet_name, startrow=0)
        summary_worksheet = writer.sheets[summary_sheet_name]
        summary_worksheet.set_column(0, 0, 50)
        summary_worksheet.write(0, 0, "Test Case Overview", header_format)


def process_file(uploaded_file, reference_text, output_directory, base_name=None):
    file_name = uploaded_file.name if hasattr(uploaded_file, 'name') else base_name
    base_name = base_name or os.path.splitext(file_name)[0]
    headline_text = f"Generating Test Scripts for {base_name}"
    placeholder = st.empty()
    scrolling_headline(True, placeholder, headline_text)

    uploaded_text = ""
    if file_name.endswith(".pdf"):
        uploaded_text += read_pdf([uploaded_file])
    elif file_name.endswith(".txt"):
        uploaded_text += read_txt([uploaded_file])
    elif file_name.endswith(".docx"):
        uploaded_text += read_docx([uploaded_file])

    if not uploaded_text:
        st.warning(f"⚠️ Skipped empty or unsupported file: {file_name}")
        return None

    combined_text = uploaded_text + reference_text
    understand_tone_and_language(combined_text)
    generative_model = initialize_generative_model()
    response = generate_script(uploaded_text, generative_model)
    df = test_cases_to_excel(response)

    testscript = {}
    for j in df["Description"]:
        testscript[j] = script(uploaded_text, j)

    excel_path = os.path.join(output_directory, f"{base_name}.xlsx")
    end_text = f"Test Script for {base_name} completed successfully!"
    scrolling_headline(False, placeholder, end_text)
    convert_to_excel(testscript, excel_path)
    return excel_path


def main():
    # Page configuration
    st.set_page_config(
        page_title="SAP Test Script Generator",
        page_icon="🚀",
        layout="wide",
        initial_sidebar_state="collapsed"
    )
    
    # Main title
    st.markdown("<div class='center-title'>🚀 SAP Test Script Generator</div>", unsafe_allow_html=True)
    
    # Add some spacing
    st.markdown("<div class='custom-spacing'></div>", unsafe_allow_html=True)

    # Input method selection with improved styling
    st.markdown("### 📁 Choose Your Input Method")
    input_method = st.radio(
        "",
        ("📄 Upload Individual Files", "📦 Upload ZIP File"),
        help="Select how you want to upload your files for processing"
    )
    
    # Directory setup
    references_directory = "References"
    output_directory = "Generated Excels"
    os.makedirs(output_directory, exist_ok=True)

    # Load reference text
    reference_text = ""
    if os.path.exists(references_directory):
        reference_text += read_pdf_from_directory(references_directory)
        reference_text += read_txt_from_directory(references_directory)
        reference_text += read_docx_from_directory(references_directory)
        
        if reference_text:
            understand_tone_and_language(reference_text)
            st.info("📚 Reference files loaded successfully!")

    if input_method == "📄 Upload Individual Files":
        st.markdown("### 📤 Upload Your Files")
        st.markdown("*Supported formats: PDF, TXT, DOCX*")
        
        uploaded_files = st.file_uploader(
            "", 
            type=['pdf', 'txt', 'docx'], 
            accept_multiple_files=True,
            help="You can upload multiple files at once"
        )
        
        if uploaded_files:
            # Display file counter
            st.markdown(f"""
            <div class='file-counter'>
                📊 {len(uploaded_files)} file(s) selected for processing
            </div>
            """, unsafe_allow_html=True)
            
            # List uploaded files
            with st.expander("📋 View Selected Files", expanded=False):
                for i, file in enumerate(uploaded_files, 1):
                    st.write(f"{i}. {file.name} ({file.size} bytes)")
            
            # Generate button
            if st.button("🚀 Generate Test Scripts"):
                with tempfile.TemporaryDirectory() as temp_dir:
                    if len(uploaded_files) == 1:
                        uploaded_file = uploaded_files[0]
                        file_base_name = os.path.splitext(uploaded_file.name)[0]
                        excel_path = process_file(uploaded_file, reference_text, temp_dir, file_base_name)
                        if excel_path:
                            with open(excel_path, 'rb') as f:
                                file_data = f.read()
                                b64 = base64.b64encode(file_data).decode()
                            st.markdown(
                                f"""
                                <a href="data:application/vnd.openxmlformats-officedocument.spreadsheetml.sheet;base64,{b64}" 
                                download="{os.path.basename(excel_path)}" class="download-button">
                                📥 Download {os.path.basename(excel_path)}
                                </a>
                                """,
                                unsafe_allow_html=True
                            )
                        else:
                            st.error("❌ Failed to generate the test script.")
                    else:
                        # Progress tracking for multiple files
                        progress_bar = st.progress(0)
                        status_text = st.empty()
                        
                        generated_excel_paths = []
                        for idx, uploaded_file in enumerate(uploaded_files):
                            status_text.text(f"Processing file {idx + 1} of {len(uploaded_files)}: {uploaded_file.name}")
                            file_base_name = os.path.splitext(uploaded_file.name)[0]
                            excel_path = process_file(uploaded_file, reference_text, temp_dir, file_base_name)
                            if excel_path:
                                generated_excel_paths.append(excel_path)
                            progress_bar.progress((idx + 1) / len(uploaded_files))

                        status_text.empty()
                        progress_bar.empty()

                        if generated_excel_paths:
                            zip_file_path = os.path.join(temp_dir, "Generated_Test_Scripts.zip")
                            with zipfile.ZipFile(zip_file_path, 'w') as zipf:
                                for file_path in generated_excel_paths:
                                    zipf.write(file_path, os.path.basename(file_path))
                            
                            with open(zip_file_path, 'rb') as f:
                                zip_data = f.read()
                                zip_b64 = base64.b64encode(zip_data).decode()
                            
                            st.success(f"✅ Successfully generated {len(generated_excel_paths)} test scripts!")
                            st.markdown(
                                f"""
                                <a href="data:application/zip;base64,{zip_b64}" 
                                download="Generated_Test_Scripts.zip" class="download-button">
                                📥 Download All Generated Test Scripts
                                </a>
                                """,
                                unsafe_allow_html=True
                            )
                        else:
                            st.error("❌ No test scripts were generated.")

    elif input_method == "📦 Upload ZIP File":
        st.markdown("### 📦 Upload ZIP File")
        st.markdown("*Upload a ZIP file containing your PDF, TXT, or DOCX files*")
        
        uploaded_zip = st.file_uploader("", type='zip')
        
        if uploaded_zip:
            st.markdown(f"""
            <div class='file-counter'>
                📦 ZIP file ready: {uploaded_zip.name} ({uploaded_zip.size} bytes)
            </div>
            """, unsafe_allow_html=True)
            
            if st.button("🚀 Generate Test Scripts"):
                with tempfile.TemporaryDirectory() as temp_dir:
                    # Extract ZIP file
                    with zipfile.ZipFile(uploaded_zip, 'r') as zip_ref:
                        zip_ref.extractall(temp_dir)
                        extracted_files = zip_ref.namelist()
                    
                    st.info(f"📂 Extracted {len(extracted_files)} files from ZIP")

                    # Find supported files
                    supported_extensions = ('*.pdf', '*.txt', '*.docx')
                    uploaded_files = []
                    for ext in supported_extensions:
                        uploaded_files.extend(glob.glob(os.path.join(temp_dir, ext)))

                    if uploaded_files:
                        st.success(f"🔍 Found {len(uploaded_files)} supported files")
                        
                        # Progress tracking
                        progress_bar = st.progress(0)
                        status_text = st.empty()
                        
                        generated_excel_paths = []
                        for idx, uploaded_file in enumerate(uploaded_files):
                            file_base_name = os.path.splitext(os.path.basename(uploaded_file))[0]
                            status_text.text(f"Processing {idx + 1} of {len(uploaded_files)}: {file_base_name}")
                            
                            with open(uploaded_file, 'rb') as f:
                                excel_path = process_file(f, reference_text, output_directory, file_base_name)
                                if excel_path:
                                    generated_excel_paths.append(excel_path)
                            progress_bar.progress((idx + 1) / len(uploaded_files))
                        
                        status_text.empty()
                        progress_bar.empty()
                        
                        if generated_excel_paths:
                            original_zip_name = os.path.splitext(uploaded_zip.name)[0]
                            zip_file_path = os.path.join(temp_dir, f"{original_zip_name}_Generated_Test_Scripts.zip")
                            
                            with zipfile.ZipFile(zip_file_path, 'w') as zipf:
                                for file_path in generated_excel_paths:
                                    zipf.write(file_path, os.path.basename(file_path))
                            
                            with open(zip_file_path, 'rb') as f:
                                zip_data = f.read()
                                zip_b64 = base64.b64encode(zip_data).decode()
                            
                            st.success(f"✅ Successfully generated {len(generated_excel_paths)} test scripts!")
                            st.markdown(
                                f"""
                                <a href="data:application/zip;base64,{zip_b64}" 
                                download="{original_zip_name}_Generated_Test_Scripts.zip" class="download-button">
                                📥 Download All Generated Test Scripts
                                </a>
                                """,
                                unsafe_allow_html=True
                            )
                        else:
                            st.error("❌ No test scripts were generated.")
                    else:
                        st.warning("⚠️ No supported files (PDF, TXT, DOCX) found in the ZIP file.")

    # Footer
    st.markdown("<div class='custom-spacing'></div>", unsafe_allow_html=True)
    st.markdown("---")
    st.markdown(
        """
        <div style='text-align: center; color: #7f8c8d; padding: 1rem;'>
            <p>🤖 Powered by AI | Built with Streamlit | SAP Test Script Generator v2.0</p>
        </div>
        """,
        unsafe_allow_html=True
    )


if __name__ == "__main__":
    main()